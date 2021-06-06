import argparse

from docx import Document
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE

from .parse import read


def generate(in_file: str, out_file: str) -> None:
    """Read the input specification from the in_file file, and generate a Word document at out_file."""
    document = Document()
    page_style = document.styles.add_style("PagePara", WD_STYLE_TYPE.PARAGRAPH)  # noqa(F841)
    panel_style = document.styles.add_style("PanelPara", WD_STYLE_TYPE.PARAGRAPH)
    panel_style.paragraph_format.left_indent = Inches(0.5)
    quote_style = document.styles.add_style("QuotePara", WD_STYLE_TYPE.PARAGRAPH)
    quote_style.paragraph_format.left_indent = Inches(1.0)
    with open(in_file, "rt") as f:
        for page, panels in read(f):
            document.add_paragraph(f"Page {page}", style="PagePara")
            for panel in range(panels):
                document.add_paragraph(f"Panel {panel+1}", style="PanelPara")
                document.add_paragraph("", style="QuotePara")
    document.save(out_file)


def main():
    parser = argparse.ArgumentParser(
        description="Generate a Word Document from page/panel specification"
    )
    parser.add_argument("--in-file", help="input file name")
    parser.add_argument("--out-file", help="output file name")
    args = parser.parse_args()
    generate(args.in_file, args.out_file)
